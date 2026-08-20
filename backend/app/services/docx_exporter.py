import io
import re
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_resume_docx(resume_data: Dict[str, Any]) -> bytes:
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    personal = resume_data.get("personal_info", {})
    name = personal.get("name") or "CANDIDATE NAME"
    role = personal.get("target_role") or "SOFTWARE ENGINEER"
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    location = personal.get("location", "")
    
    # Header: Name
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_name = title_p.add_run(name.upper())
    r_name.font.name = "Calibri"
    r_name.font.size = Pt(20)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(30, 41, 59)

    # Subtitle: Role & Contacts
    contact_parts = [p for p in [role, email, phone, location] if p]
    if contact_parts:
        c_p = doc.add_paragraph()
        r_c = c_p.add_run(" | ".join(contact_parts))
        r_c.font.name = "Calibri"
        r_c.font.size = Pt(10)
        r_c.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_heading(title: str):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(2)
        r = h.add_run(title.upper())
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)

    # Summary
    summary = resume_data.get("summary")
    if summary:
        add_heading("Professional Summary")
        p_sum = doc.add_paragraph()
        p_sum.paragraph_format.space_after = Pt(6)
        r_sum = p_sum.add_run(summary)
        r_sum.font.name = "Calibri"
        r_sum.font.size = Pt(10.5)

    # Skills
    skills = resume_data.get("skills", [])
    if skills:
        add_heading("Core Technical Skills")
        p_sk = doc.add_paragraph()
        p_sk.paragraph_format.space_after = Pt(6)
        r_sk = p_sk.add_run(" • ".join(skills))
        r_sk.font.name = "Calibri"
        r_sk.font.size = Pt(10.5)
        r_sk.font.bold = True

    # Experience
    experience = resume_data.get("experience", [])
    if experience:
        add_heading("Work Experience")
        for exp in experience:
            p_exp = doc.add_paragraph()
            p_exp.paragraph_format.space_after = Pt(4)
            if isinstance(exp, dict):
                role_str = exp.get("role") or exp.get("title") or "Software Engineer"
                comp_str = exp.get("company") or ""
                dur_str = exp.get("duration") or exp.get("dates") or ""
                desc_str = exp.get("description") or ""

                r_role = p_exp.add_run(f"{role_str}")
                r_role.font.bold = True
                r_role.font.name = "Calibri"
                r_role.font.size = Pt(11)

                if comp_str:
                    r_comp = p_exp.add_run(f" — {comp_str}")
                    r_comp.font.name = "Calibri"
                    r_comp.font.size = Pt(11)

                if dur_str:
                    r_dur = p_exp.add_run(f" ({dur_str})")
                    r_dur.font.italic = True
                    r_dur.font.size = Pt(10)

                if desc_str:
                    p_desc = doc.add_paragraph(style='List Bullet')
                    p_desc.paragraph_format.space_after = Pt(3)
                    r_desc = p_desc.add_run(desc_str)
                    r_desc.font.name = "Calibri"
                    r_desc.font.size = Pt(10)
            elif isinstance(exp, str):
                p_desc = doc.add_paragraph(style='List Bullet')
                p_desc.paragraph_format.space_after = Pt(3)
                r_desc = p_desc.add_run(exp)
                r_desc.font.name = "Calibri"
                r_desc.font.size = Pt(10)

    # Projects
    projects = resume_data.get("projects", [])
    if projects:
        add_heading("Key Projects")
        for proj in projects:
            p_proj = doc.add_paragraph()
            p_proj.paragraph_format.space_after = Pt(3)
            if isinstance(proj, dict):
                p_name = proj.get("name") or "Project"
                p_tech = proj.get("technologies") or []
                p_desc = proj.get("description") or ""

                r_pname = p_proj.add_run(f"{p_name}")
                r_pname.font.bold = True
                r_pname.font.size = Pt(11)
                
                if p_tech:
                    tech_s = ", ".join(p_tech) if isinstance(p_tech, list) else str(p_tech)
                    r_ptech = p_proj.add_run(f" [{tech_s}]")
                    r_ptech.font.size = Pt(9.5)
                    r_ptech.font.color.rgb = RGBColor(79, 70, 229)

                if p_desc:
                    p_pdesc = doc.add_paragraph(style='List Bullet')
                    p_pdesc.paragraph_format.space_after = Pt(3)
                    r_pdesc = p_pdesc.add_run(p_desc)
                    r_pdesc.font.size = Pt(10)
            elif isinstance(proj, str):
                p_pdesc = doc.add_paragraph(style='List Bullet')
                p_pdesc.paragraph_format.space_after = Pt(3)
                r_pdesc = p_pdesc.add_run(proj)
                r_pdesc.font.size = Pt(10)

    # Certifications & Achievements
    certs = resume_data.get("certifications", [])
    achievements = resume_data.get("achievements", [])
    combined_certs = certs + achievements
    if combined_certs:
        add_heading("Certifications & Achievements")
        p_c = doc.add_paragraph(style='List Bullet')
        p_c.paragraph_format.space_after = Pt(4)
        r_c = p_c.add_run(" • ".join(combined_certs))
        r_c.font.name = "Calibri"
        r_c.font.size = Pt(10)

    # Links & Portfolio
    links = resume_data.get("links", [])
    if links:
        add_heading("Links & Portfolio")
        p_l = doc.add_paragraph()
        p_l.paragraph_format.space_after = Pt(4)
        r_l = p_l.add_run(" | ".join(links))
        r_l.font.name = "Calibri"
        r_l.font.size = Pt(9.5)
        r_l.font.color.rgb = RGBColor(79, 70, 229)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
